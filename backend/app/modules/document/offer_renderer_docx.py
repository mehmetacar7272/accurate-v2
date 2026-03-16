
from __future__ import annotations

import copy
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.modules.offer.services import get_offer_document_context

OUTPUT_DIR = Path(__file__).resolve().parents[4] / "generated_offers"
ACCENT = RGBColor(122, 21, 15)
MUTED = RGBColor(71, 84, 103)
LIGHT_FILL = "FCE7E4"
HEADER_FILL = "EEF2F7"
BORDER = "D0D5DD"


def _slug(value: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "_", value or "offer")
    return normalized.strip("_") or "offer"


def _currency_symbol(code: str | None) -> str:
    value = (code or "TRY").upper()
    return {"TRY": "₺", "TL": "₺", "USD": "$", "EUR": "€"}.get(value, value)


def _money(value: Any, currency: str | None = "TRY") -> str:
    symbol = _currency_symbol(currency)
    try:
        number = f"{float(value or 0):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return f"{number} {symbol}"
    except Exception:
        return f"0,00 {symbol}"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_table_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def _set_paragraph(paragraph, *, size=10.5, bold=False, color: RGBColor | None = None, align=None, space_after=0, space_before=0):
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color


def _add_text(paragraph, text: str, *, size=10.5, bold=False, color: RGBColor | None = None):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def _add_hyperlink(paragraph, url: str, text: str, color="7A150F", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(rPr)
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _insert_page_field(paragraph, field_name: str):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fld_char1)
    r_element.append(instr)
    r_element.append(fld_char2)
    r_element.append(text)
    r_element.append(fld_char3)
    run.font.name = "Arial"
    run.font.size = Pt(7.5)


def _footer(section, context: dict[str, Any]):
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.paragraph_format.space_after = Pt(0)
    border = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '7A150F')
    border.append(top)
    footer_para._p.get_or_add_pPr().append(border)

    tbl = footer.add_table(rows=3, cols=3, width=Cm(16.4))
    tbl.autofit = False
    widths = [Cm(6.5), Cm(5.0), Cm(4.9)]
    for col, w in zip(tbl.columns, widths):
        for cell in col.cells:
            cell.width = w
    lines = [
        [context['template']['footer_company_text'], 'e-mail: info@accuratetest.com.tr', 'Tel: 0216 594 19 47'],
        [context['template'].get('footer_address_text',''), 'web: www.accuratetest.com.tr', 'Fax: 0216 594 19 17'],
        [context['template']['footer_form_full_text'], '', ''],
    ]
    for r_idx, row in enumerate(lines):
        for c_idx, value in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            _set_table_cell_margins(cell, 0, 0, 0, 0)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx < 2 else WD_ALIGN_PARAGRAPH.RIGHT
            _add_text(p, value, size=7.5, color=MUTED)
            p.paragraph_format.space_after = Pt(0)
    page_cell = tbl.cell(2, 2)
    p = page_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _insert_page_field(p, 'PAGE')
    _add_text(p, '/', size=7.5, color=MUTED)
    _insert_page_field(p, 'NUMPAGES')


def _add_table_header_row(table, labels, fill=HEADER_FILL, font_size=9.2):
    row = table.rows[0]
    for i, label in enumerate(labels):
        cell = row.cells[i]
        _set_cell_shading(cell, fill)
        _set_table_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        _add_text(p, label, size=font_size, bold=True, color=ACCENT)


def _style_table(table, font_size=9.5):
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_table_cell_margins(cell)
            for p in cell.paragraphs:
                _set_paragraph(p, size=font_size)


def _add_logo_header(document: Document, context: dict[str, Any]):
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Cm(4.6), Cm(8.2), Cm(3.6)]
    for col, w in zip(table.columns, widths):
        for cell in col.cells:
            cell.width = w
    left = table.cell(0,0)
    title = table.cell(0,1)
    right = table.cell(0,2)
    for cell in (left, title, right):
        _set_table_cell_margins(cell, 0, 0, 0, 0)
    company_logo = context['template'].get('company_logo_path')
    accred_logo = context['template'].get('accreditation_logo_path')
    from_path = lambda p: p and (Path(p) if Path(p).exists() else Path(__file__).resolve().parents[4] / p)
    lp = from_path(company_logo)
    rp = from_path(accred_logo)
    if lp and lp.exists():
        left.paragraphs[0].add_run().add_picture(str(lp), width=Cm(3.6))
    if rp and rp.exists():
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.paragraphs[0].add_run().add_picture(str(rp), width=Cm(2.0))
    p = title.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text(p, context['template']['document_title'], size=15, bold=True, color=ACCENT)
    p.paragraph_format.space_after = Pt(2)


def _meta_table(document: Document, context: dict[str, Any]):
    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    for col, w in zip(table.columns, [Cm(3.2), Cm(13.2)]):
        for cell in col.cells:
            cell.width = w
    rows = [("Teklif No", context['offer_no']), ("Tarih", context['date'])]
    for idx, (label, value) in enumerate(rows):
        left = table.cell(idx, 0)
        right = table.cell(idx, 1)
        _set_cell_shading(left, LIGHT_FILL)
        for cell in (left, right):
            _set_table_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _add_text(left.paragraphs[0], label, size=10.3, bold=True, color=ACCENT)
        _add_text(right.paragraphs[0], value, size=10.3)


def _section_heading(document: Document, text: str):
    p = document.add_paragraph()
    _add_text(p, text, size=12, bold=True, color=ACCENT)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)


def _scope_table(document: Document, section: dict[str, Any]):
    rows = section.get('summary_rows') or []
    if not rows:
        return
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    for col, w in zip(table.columns, [Cm(6.0), Cm(10.0)]):
        for cell in col.cells:
            cell.width = w
    _add_table_header_row(table, ['Kapsam Bilgileri', 'Değer'])
    for item in rows:
        row = table.add_row()
        _add_text(row.cells[0].paragraphs[0], str(item.get('label') or '-'), size=9.5)
        _add_text(row.cells[1].paragraphs[0], str(item.get('value') or '-'), size=9.5)
    _style_table(table, 9.4)


def _tests_price_table(document: Document, section: dict[str, Any], currency: str):
    outer = document.add_table(rows=2, cols=2)
    outer.autofit = False
    for col, w in zip(outer.columns, [Cm(10.1), Cm(5.9)]):
        for cell in col.cells:
            cell.width = w
    _add_table_header_row(outer, ['Uygulanacak Testler', 'Fiyat Özeti'])
    tests_cell = outer.cell(1,0)
    price_cell = outer.cell(1,1)
    tests_para = tests_cell.paragraphs[0]
    tests_para.paragraph_format.space_after = Pt(0)
    tests = section.get('tests') or []
    if tests:
        for i, test in enumerate(tests):
            p = tests_para if i == 0 else tests_cell.add_paragraph()
            _add_text(p, f"• {test['test_name']}", size=9.5)
            p.paragraph_format.space_after = Pt(0)
    else:
        _add_text(tests_para, '-', size=9.5)
    rows = [
        ('Hizmet Bedeli', _money(section.get('service_price'), currency)),
        ('Yol ve Konaklama', _money(section.get('travel_price'), currency)),
        ('Raporlama', _money(section.get('report_price'), currency)),
        ('Ara Toplam', _money(section.get('subtotal'), currency)),
        (f"KDV ({section.get('vat_rate_label')})", _money(section.get('vat_amount'), currency)),
        ('Genel Toplam', _money(section.get('grand_total'), currency)),
    ]
    inner = price_cell.add_table(rows=0, cols=2)
    inner.autofit = False
    for col, w in zip(inner.columns, [Cm(3.6), Cm(2.1)]):
        for cell in col.cells:
            cell.width = w
    for idx, (label, value) in enumerate(rows):
        row = inner.add_row()
        p1 = row.cells[0].paragraphs[0]
        p2 = row.cells[1].paragraphs[0]
        _add_text(p1, label, size=9.3, bold=(idx == len(rows)-1))
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_text(p2, value, size=9.3, bold=(idx == len(rows)-1))
        _set_table_cell_margins(row.cells[0], 0, 0, 0, 0)
        _set_table_cell_margins(row.cells[1], 0, 40, 0, 80)
    _style_table(outer, 9.3)


def _summary_table(document: Document, context: dict[str, Any]):
    rows = context.get('summary_rows') or []
    table = document.add_table(rows=1, cols=7)
    table.autofit = False
    widths = [Cm(1.7), Cm(5.2), Cm(2.2), Cm(2.5), Cm(1.8), Cm(2.4), Cm(2.6)]
    for col, w in zip(table.columns, widths):
        for cell in col.cells:
            cell.width = w
    _add_table_header_row(table, ['Madde No', 'Muayene Türü', 'Tahmini Gün', 'Ara Toplam', 'KDV %', 'KDV Tutarı', 'Genel Toplam'])
    for item in rows:
        row = table.add_row()
        vals = [item['row_no'], item['inspection_type_name'], str(item.get('estimated_days') or '-'), _money(item['subtotal'], context['currency']), item['vat_rate_label'], _money(item['vat_amount'], context['currency']), _money(item['grand_total'], context['currency'])]
        for i, val in enumerate(vals):
            p = row.cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0,2,4) else WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.RIGHT
            _add_text(p, str(val), size=9.0)
    totals = rows[-1]['document_totals'] if rows else {'subtotal':0,'vat_amount':0,'grand_total':0}
    for label, value in [('Toplam Ara Toplam', totals['subtotal']), ('Toplam KDV', totals['vat_amount']), ('GENEL TOPLAM', totals['grand_total'])]:
        row = table.add_row()
        row.cells[1].merge(row.cells[5])
        _add_text(row.cells[1].paragraphs[0], label, size=9.0, bold=True)
        row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_text(row.cells[6].paragraphs[0], _money(value, context['currency']), size=9.0, bold=True)
    _style_table(table, 8.9)


def _revision_table(document: Document, context: dict[str, Any]):
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    for col, w in zip(table.columns, [Cm(2.6), Cm(3.4), Cm(10.0)]):
        for cell in col.cells:
            cell.width = w
    _add_table_header_row(table, ['Revizyon No', 'Revizyon Tarihi', 'Açıklama'])
    for item in context.get('revision_logs') or []:
        row = table.add_row()
        vals = [item['revision_no'], item['revision_date'], item['revision_reason']]
        for i, val in enumerate(vals):
            p = row.cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            _add_text(p, str(val), size=9.0)
    _style_table(table, 9.0)


def _signature_tables(document: Document, context: dict[str, Any]):
    title_tbl = document.add_table(rows=1, cols=2)
    title_tbl.autofit = False
    for col in title_tbl.columns:
        for cell in col.cells:
            cell.width = Cm(8.2)
    _set_cell_shading(title_tbl.cell(0,0), 'FDF1EF')
    _set_cell_shading(title_tbl.cell(0,1), 'FDF1EF')
    for idx, text in enumerate([context['template']['accurate_sign_title'], context['template']['customer_sign_title']]):
        p = title_tbl.cell(0,idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_text(p, text, size=9.2, bold=True)
    body = document.add_table(rows=1, cols=2)
    body.autofit = False
    for col in body.columns:
        for cell in col.cells:
            cell.width = Cm(8.2)
    left = body.cell(0,0)
    right = body.cell(0,1)
    _add_text(left.paragraphs[0], f"YETKİLİ: {context.get('authorized_person_name') or '-'}", size=9.2)
    p = left.add_paragraph()
    _add_text(p, 'KAŞE / İMZA', size=8.8, color=MUTED)
    sig_path = context['template'].get('signature_image_path')
    if sig_path:
        path = Path(sig_path) if Path(sig_path).exists() else Path(__file__).resolve().parents[4] / sig_path
        if path.exists():
            left.add_paragraph().add_run().add_picture(str(path), width=Cm(4.8))
    _add_text(right.paragraphs[0], 'YETKİLİ:', size=9.2)
    p = right.add_paragraph()
    _add_text(p, 'KAŞE / İMZA', size=8.8, color=MUTED)


def generate_offer_docx(offer_id: int) -> Path:
    context = get_offer_document_context(offer_id)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{_slug(context['offer_no'])}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = OUTPUT_DIR / filename
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    _footer(section, context)

    _add_logo_header(document, context)
    document.add_paragraph()
    _meta_table(document, context)

    _section_heading(document, '1. TEKLİF KAPSAMI')
    p = document.add_paragraph()
    _add_text(p, context['customer_name'], size=10.4, bold=True)
    _add_text(p, ' kuruluşunun ilgili alanlarının, seçili muayene türleri kapsamında belirtilen standartlara göre muayene faaliyetlerini kapsar.', size=10.4)

    _section_heading(document, '2. REFERANS ALINAN STANDARTLAR')
    p = document.add_paragraph()
    _add_text(p, context['template']['standards_intro_text'], size=10.0)
    for link in context.get('reference_links') or []:
        p = document.add_paragraph()
        _add_hyperlink(p, link['url'], link['label'])

    _section_heading(document, '3. HİZMET BEDELLERİ VE ÖDEMELER')
    for section_doc in context.get('section_documents') or []:
        p = document.add_paragraph()
        _add_text(p, f"{section_doc['row_no']} {section_doc['inspection_type_name']}", size=11.5, bold=True, color=ACCENT)
        _scope_table(document, section_doc)
        document.add_paragraph()
        _tests_price_table(document, section_doc, context['currency'])
        document.add_paragraph()

    _section_heading(document, 'GENEL TEKLİF ÖZETİ')
    _summary_table(document, context)

    for heading, body in [
        ('3.1. Ödemeler', context['template']['payment_terms_text']),
        ('3.2. Not', context['template']['note_text']),
        ('3.3. Tahmini Süre', re.sub(r'<[^>]+>', '', context['template']['estimated_duration_text'])),
        ('3.4. Opsiyon', context['template']['option_validity_text']),
    ]:
        _section_heading(document, heading)
        p = document.add_paragraph()
        _add_text(p, body, size=10.0)

    _section_heading(document, '3.5. Diğer Şartlar')
    for chunk in (context['template']['other_terms_text'] or '').split('\n\n'):
        p = document.add_paragraph()
        if 'http://www.accuratetest.com.tr/hizmetsartlariformu.pdf' in chunk:
            before, url_after = chunk.split('http://www.accuratetest.com.tr/hizmetsartlariformu.pdf', 1)
            _add_text(p, before, size=10.0)
            _add_hyperlink(p, 'http://www.accuratetest.com.tr/hizmetsartlariformu.pdf', 'Hizmet Şartları Formu (FR.059)')
            _add_text(p, url_after.replace('http://www.accuratetest.com.tr/hizmetsartlariformu.pdf', ''), size=10.0)
        else:
            _add_text(p, chunk, size=10.0)

    _section_heading(document, 'REVİZYON GEÇMİŞİ')
    _revision_table(document, context)
    document.add_paragraph()
    _signature_tables(document, context)
    document.save(output_path)
    return output_path
